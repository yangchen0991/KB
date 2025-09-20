"""
文档处理异步任务
"""

import os
import logging
from typing import Optional
from datetime import datetime, timedelta

from celery import shared_task
from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone

logger = logging.getLogger(__name__)


@shared_task(bind=True, max_retries=3)
def process_document_async(self, document_id: int):
    """
    异步处理文档
    - 提取文本内容
    - 生成缩略图
    - OCR识别（如果是图片）
    - 更新处理状态
    """
    from .models import Document
    
    try:
        document = Document.objects.get(id=document_id)
        logger.info(f"开始处理文档: {document.title} (ID: {document_id})")
        
        # 更新状态为处理中
        document.status = 'processing'
        document.processing_progress = 0
        document.save(update_fields=['status', 'processing_progress'])
        
        # 根据文件类型进行处理
        file_ext = document.file_type.lower()
        
        if file_ext == 'pdf':
            _process_pdf(document)
        elif file_ext in ['doc', 'docx']:
            _process_word_document(document)
        elif file_ext == 'txt':
            _process_text_file(document)
        elif file_ext in ['jpg', 'jpeg', 'png', 'tiff']:
            _process_image(document)
        
        # 处理完成
        document.status = 'completed'
        document.processing_progress = 100
        document.save(update_fields=['status', 'processing_progress'])
        
        logger.info(f"文档处理完成: {document.title}")
        
        # 触发搜索索引更新
        try:
            update_search_index.delay(document_id)
        except Exception as e:
            logger.warning(f"搜索索引更新失败: {e}")
        
        return f"文档 {document.title} 处理完成"
        
    except Document.DoesNotExist:
        logger.error(f"文档不存在: ID {document_id}")
        raise
    except Exception as exc:
        logger.error(f"文档处理失败: {exc}")
        
        # 更新失败状态
        try:
            document = Document.objects.get(id=document_id)
            document.status = 'failed'
            document.save(update_fields=['status'])
        except:
            pass
        
        # 重试机制
        if self.request.retries < self.max_retries:
            logger.info(f"重试处理文档 {document_id}, 第 {self.request.retries + 1} 次")
            raise self.retry(countdown=60 * (2 ** self.request.retries))
        
        raise


def _process_pdf(document):
    """处理PDF文件"""
    try:
        import PyPDF2
        import pdfplumber
        from pdf2image import convert_from_path
        from PIL import Image
        import io
        
        # 更新进度
        document.processing_progress = 10
        document.save(update_fields=['processing_progress'])
        
        # 提取文本内容
        text_content = ""
        page_count = 0
        
        # 使用 pdfplumber 提取文本（更准确）
        with pdfplumber.open(document.file.path) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                text_content += page.extract_text() or ""
                # 更新进度
                progress = 10 + (i + 1) / page_count * 60
                document.processing_progress = int(progress)
                document.save(update_fields=['processing_progress'])
        
        document.ocr_text = text_content
        document.page_count = page_count
        document.processing_progress = 70
        document.save(update_fields=['ocr_text', 'page_count', 'processing_progress'])
        
        # 生成缩略图（第一页）
        try:
            images = convert_from_path(document.file.path, first_page=1, last_page=1, dpi=150)
            if images:
                img = images[0]
                img.thumbnail((300, 400), Image.Resampling.LANCZOS)
                
                # 保存缩略图
                thumb_io = io.BytesIO()
                img.save(thumb_io, format='JPEG', quality=85)
                thumb_file = ContentFile(thumb_io.getvalue())
                
                document.thumbnail.save(
                    f"thumb_{document.id}.jpg",
                    thumb_file,
                    save=False
                )
        except Exception as e:
            logger.warning(f"PDF缩略图生成失败: {e}")
        
        document.processing_progress = 90
        document.save(update_fields=['thumbnail', 'processing_progress'])
        
    except Exception as e:
        logger.error(f"PDF处理失败: {e}")
        raise


def _process_word_document(document):
    """处理Word文档"""
    try:
        from docx import Document as DocxDocument
        
        document.processing_progress = 20
        document.save(update_fields=['processing_progress'])
        
        # 提取文本内容
        doc = DocxDocument(document.file.path)
        text_content = ""
        
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        document.ocr_text = text_content
        document.page_count = len(doc.paragraphs)
        document.processing_progress = 90
        document.save(update_fields=['ocr_text', 'page_count', 'processing_progress'])
        
    except Exception as e:
        logger.error(f"Word文档处理失败: {e}")
        raise


def _process_text_file(document):
    """处理文本文件"""
    try:
        document.processing_progress = 20
        document.save(update_fields=['processing_progress'])
        
        # 读取文本内容
        with open(document.file.path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        document.ocr_text = text_content
        document.page_count = 1
        document.processing_progress = 90
        document.save(update_fields=['ocr_text', 'page_count', 'processing_progress'])
        
    except UnicodeDecodeError:
        # 尝试其他编码
        try:
            with open(document.file.path, 'r', encoding='gbk') as f:
                text_content = f.read()
            document.ocr_text = text_content
            document.page_count = 1
            document.processing_progress = 90
            document.save(update_fields=['ocr_text', 'page_count', 'processing_progress'])
        except Exception as e:
            logger.error(f"文本文件处理失败: {e}")
            raise
    except Exception as e:
        logger.error(f"文本文件处理失败: {e}")
        raise


def _process_image(document):
    """处理图片文件"""
    try:
        import pytesseract
        from PIL import Image
        import io
        
        document.processing_progress = 20
        document.save(update_fields=['processing_progress'])
        
        # OCR识别
        img = Image.open(document.file.path)
        
        # 生成缩略图
        img_copy = img.copy()
        img_copy.thumbnail((300, 400), Image.Resampling.LANCZOS)
        
        thumb_io = io.BytesIO()
        img_copy.save(thumb_io, format='JPEG', quality=85)
        thumb_file = ContentFile(thumb_io.getvalue())
        
        document.thumbnail.save(
            f"thumb_{document.id}.jpg",
            thumb_file,
            save=False
        )
        
        document.processing_progress = 50
        document.save(update_fields=['thumbnail', 'processing_progress'])
        
        # OCR文字识别
        try:
            # 支持中英文识别
            ocr_text = pytesseract.image_to_string(img, lang='chi_sim+eng')
            confidence = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            
            # 计算平均置信度
            confidences = [int(conf) for conf in confidence['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            document.ocr_text = ocr_text
            document.ocr_confidence = avg_confidence / 100.0  # 转换为0-1范围
            document.ocr_language = 'chi_sim+eng'
            
        except Exception as e:
            logger.warning(f"OCR识别失败: {e}")
            document.ocr_text = ""
            document.ocr_confidence = 0.0
        
        document.page_count = 1
        document.processing_progress = 90
        document.save(update_fields=[
            'ocr_text', 'ocr_confidence', 'ocr_language', 
            'page_count', 'processing_progress'
        ])
        
    except Exception as e:
        logger.error(f"图片处理失败: {e}")
        raise


@shared_task
def update_search_index(document_id: int):
    """更新搜索索引"""
    try:
        from .models import Document
        
        document = Document.objects.get(id=document_id)
        
        # 如果启用了 Elasticsearch
        try:
            from django_elasticsearch_dsl.registries import registry
            registry.update(document)
            logger.info(f"搜索索引更新成功: {document.title}")
        except ImportError:
            logger.info("Elasticsearch 未启用，跳过索引更新")
        
    except Document.DoesNotExist:
        logger.error(f"文档不存在: ID {document_id}")
    except Exception as e:
        logger.error(f"搜索索引更新失败: {e}")


@shared_task
def cleanup_failed_documents():
    """清理处理失败的文档"""
    from .models import Document
    from datetime import datetime, timedelta
    
    # 清理24小时前处理失败的文档
    cutoff_time = timezone.now() - timedelta(hours=24)
    failed_docs = Document.objects.filter(
        status='failed',
        updated_at__lt=cutoff_time
    )
    
    count = 0
    for doc in failed_docs:
        try:
            # 删除文件
            if doc.file and os.path.exists(doc.file.path):
                os.remove(doc.file.path)
            if doc.thumbnail and os.path.exists(doc.thumbnail.path):
                os.remove(doc.thumbnail.path)
            
            # 删除记录
            doc.delete()
            count += 1
            
        except Exception as e:
            logger.error(f"清理失败文档出错: {e}")
    
    logger.info(f"清理了 {count} 个失败的文档")
    return f"清理了 {count} 个失败的文档"


@shared_task
def generate_document_statistics():
    """生成文档统计报告"""
    from .models import Document, Category, Tag
    from django.db.models import Count, Sum, Avg
    
    try:
        # 总体统计
        total_docs = Document.objects.count()
        total_size = Document.objects.aggregate(Sum('file_size'))['file_size__sum'] or 0
        avg_size = Document.objects.aggregate(Avg('file_size'))['file_size__avg'] or 0
        
        # 按分类统计
        category_stats = Category.objects.annotate(
            doc_count=Count('documents')
        ).order_by('-doc_count')[:10]
        
        # 按文件类型统计
        type_stats = Document.objects.values('file_type').annotate(
            count=Count('id'),
            total_size=Sum('file_size')
        ).order_by('-count')[:10]
        
        # 处理状态统计
        status_stats = Document.objects.values('status').annotate(
            count=Count('id')
        )
        
        stats = {
            'total_documents': total_docs,
            'total_size': total_size,
            'average_size': avg_size,
            'by_category': list(category_stats.values()),
            'by_type': list(type_stats),
            'by_status': list(status_stats),
            'generated_at': timezone.now().isoformat()
        }
        
        logger.info(f"统计报告生成完成: {total_docs} 个文档")
        return stats
        
    except Exception as e:
        logger.error(f"统计报告生成失败: {e}")
        raise